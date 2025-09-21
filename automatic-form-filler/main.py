import docx
import re
import requests
import dotenv
import os
import json
import datetime
from send_report_gmail import send_email_with_attachment

def get_fileInfoByHash(fileSha1Hash, VT_API_KEY):
    url = f"https://www.virustotal.com/api/v3/files/{fileSha1Hash}"
    headers = {
        "accept": "application/json",
        "x-apikey": f"{VT_API_KEY}"
    }
    response = requests.get(url, headers=headers)

    try:
        response.raise_for_status() 
    except requests.exceptions.HTTPError as err:
        print(f"HTTP Error occured: {err}")
        exit()
    return response.json()

def get_sampleData():
    with open('Documents/sample-data.json','r') as f:
        return json.load(f)

def replace_placeHolders(doc, parsedData):
    for para in doc.paragraphs:
        for data in parsedData:
            if re.search(data, para.text):
                replacement = parsedData[data]
                if type(replacement) == list:
                    replacement = ', '.join(f'"{value}"' for value in replacement)
                if type(replacement) == dict:
                    replacement = '\n'.join(f'- {key}: {value}' for key, value in replacement.items())
                para.text = re.sub(data, str(replacement), para.text)

def populate_placeHolders(holders, apiResult):
    populatedData = {}
    for holder in holders:
        match holder:
            case "{{ReportGenerated}}":
                currentTime = datetime.datetime.now()
                populatedData.update({holder: str(currentTime)})
            case "{{AnalystName}}": 
                analystName = 'Horse69420'
                populatedData.update({holder: analystName})
            case "{{MD5}}":
                sampleMd5 = apiResult['data']['attributes']['md5']
                populatedData.update({holder: sampleMd5})
            case "{{SHA1}}":
                sampleSha1 = apiResult['data']['attributes']['sha1']
                populatedData.update({holder: sampleSha1})
            case "{{SHA256}}":
                sampleSha256 = apiResult['data']['attributes']['sha256']
                populatedData.update({holder: sampleSha256})
            case "{{FileSize}}":
                sampleSize = apiResult['data']['attributes']['size']
                populatedData.update({holder: sampleSize})
            case "{{FileName}}":
                sampleNames = apiResult['data']['attributes']['names']
                populatedData.update({holder: sampleNames})
            case "{{DetectionRatio}}":
                sampleStats = apiResult['data']['attributes']['last_analysis_stats']
                maliciousCount = sampleStats['malicious']
                total = sum(sampleStats[stat] for stat in sampleStats)
                populatedData.update({holder: f'{maliciousCount} out of {total} security vendors flagged the sample as "malicious".'})
            case "{{MalwareSignature}}":
                sampleSignature = apiResult['data']['attributes']['popular_threat_classification']['suggested_threat_label']
                populatedData.update({holder: sampleSignature})
            case "{{Vendors}}":
                vendorObjs = apiResult['data']['attributes']['last_analysis_results']
                vendors = {}
                for vendor in vendorObjs: 
                    if len(vendors) != 10:
                        if vendorObjs[vendor]['result'] != None:
                            vendors.update({f'{vendor}': f'{vendorObjs[vendor]["result"]}'})
                populatedData.update({holder: vendors})
            case _:
                print(f'\nNothing to do for this matched case: {holder}\n')
    return json.dumps(populatedData)

def get_placeHolders(doc):
    holders = []
    hPattern = re.compile(r'\{\{[\w]+\}\}')
    for para in doc.paragraphs:
        if hPattern.search(para.text):
            holders.extend(hPattern.findall(para.text))
    return holders

def get_env(secret): 
    script_dir = os.path.dirname(os.path.abspath(__file__))
    dotenv_path = os.path.join(script_dir, "Secrets", ".env")
    dotenv.load_dotenv(dotenv_path)    
    val = os.getenv(secret)
    if not val:
        raise RuntimeError(f"{secret} not set; Recheck .env file.")
    return val

def main():
    # Logic to fill up the template malware report.
    fileSha1Sum = "" # Input file SHA-1 hash here
    template_report = "Documents/malware_report_template.docx"
    output_filename = "Documents/filled_report.docx"
    doc = docx.Document(template_report)
    holders = get_placeHolders(doc)
    # apiResult = get_fileInfoByHash(fileSha1Sum, get_env("VIRUS_TOTAL_API_KEY")) # Fetch real data
    apiResult = get_sampleData() # Fetch sample data
    jsonData = populate_placeHolders(holders, apiResult)
    replace_placeHolders(doc, json.loads(jsonData))
    doc.save(output_filename)

    # Logic to send report via email. (Can comment out if not needed)
    sender_email = "" # Insert sender email here
    to_email = "" # Insert recipient email here
    subject = "Your Automated Malware Report"
    message_body = "Hi, here is the automated malware analysis report you requested."
    send_email_with_attachment(sender_email, to_email, subject, message_body, output_filename)

if __name__ == "__main__":
    main()

